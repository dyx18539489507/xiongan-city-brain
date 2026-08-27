from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "docs" / "实验评估报告.md"
OUTPUT = ROOT / "docs" / "实验评估报告.docx"
FIGURE_DIR = ROOT / "experiments" / "formal_2026" / "figures"

INK = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6875"
LIGHT_FILL = "F4F6F9"
TABLE_FILL = "E8EEF5"
WHITE = "FFFFFF"
GRID = "B7C2CE"
BODY_FONT = "Calibri"
BODY_EAST_ASIA = "宋体"
HEADING_EAST_ASIA = "微软雅黑"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


FIGURES = [
    ("figure_01_average_speed.png", "图 1 不同控制算法平均速度（95% CI，n=5）"),
    ("figure_02_average_queue.png", "图 2 不同控制算法平均排队（95% CI，n=5）"),
    ("figure_03_max_queue.png", "图 3 不同控制算法最大排队（95% CI，n=5）"),
    ("figure_04_fuel.png", "图 4 不同控制算法燃油消耗（95% CI，n=5）"),
    ("figure_05_realtime_factor.png", "图 5 不同控制算法仿真实时因子（95% CI，n=5）"),
]


def set_run_font(run, *, name=BODY_FONT, east_asia=BODY_EAST_ASIA, size=None,
                 color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        widths_dxa[-1] += CONTENT_WIDTH_DXA - sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_EAST_ASIA)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_EAST_ASIA)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_EAST_ASIA)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = doc.styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(9.5)
    caption.font.bold = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_EAST_ASIA)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("面向雄安新区城市大脑 · 实验评估报告")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("快速实验评估版  |  ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(fp)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026-08-20 · 快速实验评估版")
    set_run_font(run, size=9, color=MUTED)


def add_inline_runs(paragraph, text, *, size=None, color=None):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", east_asia=BODY_EAST_ASIA,
                         size=(size or 10), color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D5DEE8", size=5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(f"{label}  ")
    set_run_font(r, east_asia=HEADING_EAST_ASIA, size=10.5, color=DARK_BLUE, bold=True)
    add_inline_runs(p, text, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(86)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实验评估报告")
    set_run_font(r, east_asia=HEADING_EAST_ASIA, size=14, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("面向雄安新区“城市大脑”的\n车路云一体化协同管控算法与仿真平台")
    set_run_font(r, east_asia=HEADING_EAST_ASIA, size=25, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(46)
    r = p.add_run("快速实验评估版 · 真实运行数据")
    set_run_font(r, east_asia=HEADING_EAST_ASIA, size=12, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [2520, 6840])
    set_table_borders(meta, color=WHITE, size=0)
    values = [
        ("评估日期", "2026-08-20"),
        ("实验规模", "22 runs（20 次快速矩阵 + 2 次先导实验）"),
        ("代码版本", "Git commit 5e6e30750b499fb33837bff9f4cb18fa6e8a5e0d"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        set_cell_shading(row.cells[0], TABLE_FILL)
        for cell in row.cells:
            set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), east_asia=HEADING_EAST_ASIA, size=9.5,
                     color=DARK_BLUE, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9.5, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_callout(
        doc,
        "证据边界",
        "本报告在用户限定的 15 分钟实验窗口内形成，完成了代码审计和 20 次 60 s 多种子快速实跑。它不是原计划的多负荷长时正式矩阵；缺失数据均写为 NA。",
    )
    doc.add_page_break()


def add_contents(doc, headings):
    p = doc.add_paragraph("目录", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    for level, text in headings:
        if level != 2:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        add_inline_runs(p, text, size=10.5, color=INK)
    doc.add_page_break()


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows


def table_widths(rows):
    columns = len(rows[0])
    if columns == 8:
        weights = [0.72, 1.05, 1.0, 1.0, 0.78, 1.02, 1.05, 0.9]
    elif columns == 5:
        weights = [0.9, 1.75, 0.9, 1.0, 1.0]
    elif columns == 4:
        weights = [1.05, 2.1, 1.15, 2.2]
    else:
        weights = [max(0.75, min(2.5, max(len(r[i]) for r in rows) / 10)) for i in range(columns)]
    total = sum(weights)
    widths = [round(CONTENT_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_data_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_header(table.rows[0])
    body_size = 8.1 if len(rows[0]) >= 7 else 8.8
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, value, size=(8.8 if row_idx == 0 else body_size),
                            color=(DARK_BLUE if row_idx == 0 else INK))
            for run in p.runs:
                if row_idx == 0:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_figures(doc):
    doc.add_page_break()
    p = doc.add_paragraph("实验统计图", style="Heading 2")
    p.paragraph_format.space_before = Pt(0)
    for idx, (filename, caption) in enumerate(FIGURES, start=1):
        path = FIGURE_DIR / filename
        if not path.exists():
            continue
        image_p = doc.add_paragraph()
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.paragraph_format.space_before = Pt(2)
        image_p.paragraph_format.space_after = Pt(0)
        image_p.paragraph_format.keep_with_next = True
        run = image_p.add_run()
        run.add_picture(str(path), width=Inches(5.45))
        caption_p = doc.add_paragraph(style="Caption")
        caption_p.paragraph_format.keep_with_next = False
        caption_p.add_run(caption)
        source_p = doc.add_paragraph()
        source_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_p.paragraph_format.space_before = Pt(0)
        source_p.paragraph_format.space_after = Pt(8)
        source_p.paragraph_format.keep_with_next = False
        set_run_font(source_p.add_run("数据来源：formal_2026 真实运行 raw 数据，经统一统计脚本生成。"),
                     size=8.5, color=MUTED)


def add_markdown_body(doc, text):
    lines = text.splitlines()
    inserted_figures = False
    idx = 0
    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if not line:
            idx += 1
            continue
        if line.startswith("# "):
            idx += 1
            continue
        if line.startswith(">"):
            idx += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            idx += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            idx += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            add_data_table(doc, parse_table(table_lines))
            continue
        if re.match(r"^-\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, re.sub(r"^-\s+", "", line))
            idx += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", line))
            idx += 1
            continue
        p = doc.add_paragraph()
        if "`" in line:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline_runs(p, line)
        if "图 1--5 只展示有真实观测的数据" in line and not inserted_figures:
            add_figures(doc)
            inserted_figures = True
        idx += 1


def build():
    text = SOURCE.read_text(encoding="utf-8")
    headings = []
    for line in text.splitlines():
        stripped = line.strip()
        match = re.match(r"^(#{2,3})\s+(.+)$", stripped)
        if match:
            headings.append((len(match.group(1)), match.group(2)))

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "面向雄安新区城市大脑车路云一体化协同管控算法与仿真平台实验评估报告"
    doc.core_properties.subject = "2026-08-20 快速实验评估版"
    doc.core_properties.author = "项目实验评估组"
    doc.core_properties.keywords = "雄安新区, 城市大脑, SUMO, 云边端, 交通控制, 实验评估"

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    add_cover(doc)
    add_contents(doc, headings)
    add_markdown_body(doc, text)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
